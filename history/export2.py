import os
import re
import ssl
import urllib3
from pathlib import Path
from urllib.parse import urljoin

import requests
from tcms_api import TCMS
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
import config


# Tymczasowo wyłączamy weryfikację SSL,
# bo w Twoim środowisku certyfikat powodował błąd.
ssl._create_default_https_context = ssl._create_unverified_context

# Folder na pobrane obrazki
DOWNLOAD_DIR = Path("../downloaded_images")
DOWNLOAD_DIR.mkdir(exist_ok=True)


def connect_to_kiwi():
    """Tworzy połączenie z Kiwi TCMS przez XML-RPC."""
    return TCMS(
        config.KIWI_URL,
        username=config.KIWI_USERNAME,
        password=config.KIWI_PASSWORD
    )


def get_test_plans(tcms):
    """Pobiera wszystkie test plany."""
    return tcms.exec.TestPlan.filter({})


def show_test_plans(plans):
    """Wyświetla listę planów."""
    print("\nDostępne Test Plany:\n")
    for plan in plans:
        print(f"ID: {plan['id']} | Nazwa: {plan['name']}")


def ask_for_plan_id(plans):
    """Pyta użytkownika o ID planu i sprawdza, czy jest poprawne."""
    valid_ids = {plan["id"] for plan in plans}

    while True:
        user_input = input("\nPodaj ID test planu: ").strip()

        if not user_input.isdigit():
            print("To nie jest liczba. Spróbuj ponownie.")
            continue

        plan_id = int(user_input)

        if plan_id not in valid_ids:
            print("Nie ma takiego planu na liście. Spróbuj ponownie.")
            continue

        return plan_id


def get_plan_name(plans, plan_id):
    """Zwraca nazwę planu po ID."""
    for plan in plans:
        if plan["id"] == plan_id:
            return plan["name"]
    return f"Test Plan {plan_id}"


def get_test_cases_from_plan(tcms, plan_id):
    """Pobiera wszystkie test case z wybranego planu."""
    return tcms.exec.TestCase.filter({"plan": plan_id})

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing

def set_doc_style(document):
    """Ustawia podstawowy wygląd dokumentu."""
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def clean_inline_formatting(text):
    if not text:
        return ""

    text = str(text)

    # **tekst** -> tekst
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # usuwamy markdown obrazka z tekstu
    text = re.sub(r'!\[.*?\]\((/uploads/attachments/[^)]+)\)', '', text)

    return text.strip()

def add_images_to_container(container, raw_text, width_cm=10):
    image_paths = extract_image_paths(raw_text)

    for image_path in image_paths:
        local_file = download_image(image_path)

        if local_file and local_file.exists():
            try:
                paragraph = container.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(str(local_file), width=Cm(width_cm))
            except Exception as e:
                container.add_paragraph(f"[Nie udało się wstawić obrazka: {local_file.name}]")
                print(f"Błąd przy wstawianiu obrazka: {e}")

def add_images_to_cell(cell, raw_text):
    image_paths = extract_image_paths(raw_text)

    for image_path in image_paths:
        local_file = download_image(image_path)

        if local_file and local_file.exists():
            try:
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(str(local_file), width=Cm(10))
            except Exception as e:
                cell.add_paragraph(f"[Nie udało się wstawić obrazka: {local_file.name}]")
                print(f"Błąd przy wstawianiu obrazka do komórki: {e}")

def extract_image_paths(text):
    if not text:
        return []

    pattern = r'!\[.*?\]\((/uploads/attachments/[^)]+)\)'
    return re.findall(pattern, str(text))


def download_image(relative_path):
    """
    Pobiera obrazek z Kiwi na podstawie ścieżki względnej.
    Zwraca lokalną ścieżkę pliku lub None.
    """
    try:
        full_url = urljoin(config.KIWI_BASE_WEB_URL, relative_path)
        filename = os.path.basename(relative_path)
        local_path = DOWNLOAD_DIR / filename

        response = requests.get(full_url, verify=False, timeout=30)
        response.raise_for_status()

        with open(local_path, "wb") as f:
            f.write(response.content)

        return local_path
    except Exception as e:
        print(f"Nie udało się pobrać obrazka: {relative_path}")
        print(f"Błąd: {e}")
        return None


def parse_markdown_table(lines):
    """
    Zamienia tabelę markdownową:
    | Krok | Czynność | Oczekiwany rezultat |
    |------|----------|---------------------|
    | 1 | aaa | bbb |
    na listę wierszy.
    """
    rows = []

    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue

        # separator |----|----|
        if re.match(r"^\|[\-\s|:]+\|?$", line):
            continue

        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)

    return rows


def add_word_table(document, rows):
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # nagłówek
    header_cells = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        header_cells[i].text = clean_inline_formatting(value)

    # dane
    for row_data in rows[1:]:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if i < len(row_cells):
                cleaned_text = clean_inline_formatting(value)
                row_cells[i].text = cleaned_text

                # obrazki przypisane do tej konkretnej komórki
                add_images_to_container(row_cells[i], value, width_cm=8)

    document.add_paragraph("")


def add_bullet_paragraph(document, text):
    """Główna lista (*)"""
    p = document.add_paragraph(style="List Bullet")
    p.add_run(clean_inline_formatting(text))

    fmt = p.paragraph_format
    fmt.left_indent = Cm(0.6)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

    add_images_to_container(document, text, width_cm=12)


def add_sub_bullet_paragraph(document, text):
    """Podlista (-)"""
    p = document.add_paragraph(style="List Bullet")
    p.add_run(clean_inline_formatting(text))

    fmt = p.paragraph_format
    fmt.left_indent = Cm(1.3)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

    add_images_to_container(document, text, width_cm=12)


def add_normal_paragraph(document, text):
    cleaned = clean_inline_formatting(text)

    if cleaned:
        p = document.add_paragraph(cleaned)
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    add_images_to_container(document, text, width_cm=12)


def add_images_from_text(document, raw_text):
    """Pobiera i wstawia obrazki znalezione w treści test case."""
    image_paths = extract_image_paths(raw_text)

    if not image_paths:
        return

    document.add_paragraph("Załączone obrazy", style="Heading 3")

    for image_path in image_paths:
        local_file = download_image(image_path)

        if local_file and local_file.exists():
            try:
                document.add_picture(str(local_file), width=Cm(14))
                document.add_paragraph(local_file.name)
            except Exception as e:
                document.add_paragraph(f"Nie udało się wstawić obrazka: {local_file.name}")
                print(f"Błąd przy wstawianiu obrazka do Worda: {e}")


def add_formatted_case_text(document, raw_text):
    if not raw_text:
        document.add_paragraph("Brak treści")
        return

    lines = str(raw_text).splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # <h2>...</h2>
        h2_match = re.match(r"<h2>\s*(.*?)\s*</h2>", line, re.IGNORECASE)
        if h2_match:
            title = clean_inline_formatting(h2_match.group(1))
            document.add_paragraph(title, style="Heading 3")
            i += 1
            continue

        # tabela markdownowa
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            rows = parse_markdown_table(table_lines)
            add_word_table(document, rows)
            continue

        # lista punktowana
        if line.startswith("* "):
            add_bullet_paragraph(document, line[2:])
            i += 1
            continue

        if line.startswith("- "):
            add_sub_bullet_paragraph(document, line[2:])
            i += 1
            continue

        # zwykły tekst
        add_normal_paragraph(document, line)
        i += 1


def add_case_section(document, case):
    case_id = case.get("id", "")
    summary = case.get("summary", "Brak nazwy")
    text = case.get("text", "Brak treści")

    heading = document.add_paragraph(style="Heading 2")
    heading.add_run(f"TC-{case_id}: {summary}")

    info_table = document.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"

    row1 = info_table.rows[0].cells
    row1[0].text = "ID"
    row1[1].text = str(case_id)

    row2 = info_table.rows[1].cells
    row2[0].text = "Nazwa"
    row2[1].text = summary

    document.add_paragraph("")
    add_formatted_case_text(document, text)
    document.add_paragraph("")


def export_plan_to_docx(plan_name, plan_id, cases, output_path):
    """Tworzy dokument Word z wybranego test planu."""
    document = Document()
    set_doc_style(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f"Test Plan: {plan_name}")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(f"ID planu: {plan_id}")

    summary = document.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary.add_run(f"Liczba test case: {len(cases)}")

    document.add_paragraph("")

    for case in cases:
        add_case_section(document, case)

    document.save(output_path)


def main():
    print("Łączenie z Kiwi TCMS...")
    tcms = connect_to_kiwi()

    print("Pobieranie test planów...")
    plans = get_test_plans(tcms)

    if not plans:
        print("Nie znaleziono żadnych test planów.")
        return

    show_test_plans(plans)
    selected_plan_id = ask_for_plan_id(plans)
    plan_name = get_plan_name(plans, selected_plan_id)

    print(f"\nPobieranie test case z planu ID = {selected_plan_id}...")
    cases = get_test_cases_from_plan(tcms, selected_plan_id)

    output_file = f"TestPlan_{selected_plan_id}.docx"
    export_plan_to_docx(plan_name, selected_plan_id, cases, output_file)

    print(f"\nGotowe. Plik zapisany jako: {output_file}")


if __name__ == "__main__":
    main()