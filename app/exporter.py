from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

import config
import re
from datetime import datetime
from app.images import download_image, extract_image_paths
from app.parser import clean_inline_formatting, split_case_text


def _ensure_output_dir() -> Path:
    output_dir = Path(config.OUTPUT_DIR)
    output_dir.mkdir(exist_ok=True)
    return output_dir


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line_spacing: float = 1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def _set_doc_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    fmt = style.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def _add_images_to_container(container, raw_text: str | None, width_cm: float) -> None:
    image_paths = extract_image_paths(raw_text)

    for image_path in image_paths:
        local_file = download_image(image_path)

        if local_file and local_file.exists():
            try:
                paragraph = container.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(str(local_file), width=Cm(width_cm))
                _set_paragraph_spacing(paragraph)
            except Exception as exc:
                container.add_paragraph(f"[Nie udało się wstawić obrazka: {local_file.name}]")
                print(f"Błąd przy wstawianiu obrazka: {exc}")


def _add_paragraph(document: Document, text: str) -> None:
    cleaned = clean_inline_formatting(text)
    if cleaned:
        paragraph = document.add_paragraph(cleaned)
        _set_paragraph_spacing(paragraph)

    _add_images_to_container(document, text, width_cm=12)


def _add_bullet(document: Document, text: str) -> None:
    cleaned = clean_inline_formatting(text)
    if cleaned:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(cleaned)

        fmt = paragraph.paragraph_format
        fmt.left_indent = Cm(0.6)
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0

    _add_images_to_container(document, text, width_cm=12)


def _add_sub_bullet(document: Document, text: str) -> None:
    cleaned = clean_inline_formatting(text)

    if cleaned:
        p = document.add_paragraph(f"- {cleaned}")  # <- kreska zamiast kropki

        fmt = p.paragraph_format
        fmt.left_indent = Cm(1.3)   # wcięcie (zostaje)
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0

    _add_images_to_container(document, text, width_cm=12)


from docx.shared import Cm


def _set_column_width(table, col_idx: int, width_cm: float) -> None:
    """Ustawia szerokość całej kolumny we wszystkich wierszach."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _is_steps_table(rows: list[list[str]]) -> bool:
    """Sprawdza, czy to tabela kroków testowych."""
    if not rows:
        return False

    header = [str(cell).strip().lower() for cell in rows[0]]
    return len(header) >= 3 and header[0] == "krok"


from docx import Document
from docx.shared import Cm


def _set_column_width(table, col_idx: int, width_cm: float) -> None:
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _is_steps_table(rows: list[list[str]]) -> bool:
    if not rows:
        return False

    header = [str(cell).strip().lower() for cell in rows[0]]
    return len(header) >= 3 and header[0] == "krok"


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        header_cells[index].text = clean_inline_formatting(value)

    for row_data in rows[1:]:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_data):
            if index < len(row_cells):
                row_cells[index].text = clean_inline_formatting(value)
                _add_images_to_container(row_cells[index], value, width_cm=6.5)

    if _is_steps_table(rows) and col_count >= 3:
        # szerokość zbliżona do górnej tabeli
        widths = [1.2, 8.1, 8.1]   # razem 17.4 cm
        for col_idx, width in enumerate(widths):
            _set_column_width(table, col_idx, width)

    document.add_paragraph("")


def _add_case_content(document: Document, raw_text: str | None) -> None:
    blocks = split_case_text(raw_text)

    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            document.add_paragraph(block["text"], style="Heading 3")
        elif block_type == "paragraph":
            _add_paragraph(document, block["text"])
        elif block_type == "bullet":
            _add_bullet(document, block["text"])
        elif block_type == "sub_bullet":
            _add_sub_bullet(document, block["text"])
        elif block_type == "table":
            _add_table(document, block["rows"])


def _add_case_section(document: Document, case: dict[str, Any]) -> None:
    case_id = case.get("id", "")
    summary = case.get("summary", "Brak nazwy")
    text = case.get("text", "Brak treści")

    heading = document.add_paragraph(style="Heading 2")
    heading.add_run(f"TC-{case_id}: {summary}")

    info_table = document.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"
    info_table.autofit = False

    # ustaw szerokości kolumn
    for row in info_table.rows:
        row.cells[0].width = Cm(2.5)  # etykiety (ID, Nazwa)
        row.cells[1].width = Cm(15)  # wartości

    # wypełnienie danych
    row1 = info_table.rows[0].cells
    row1[0].text = "ID"
    row1[1].text = str(case_id)

    row2 = info_table.rows[1].cells
    row2[0].text = "Nazwa"
    row2[1].text = str(summary)

    document.add_paragraph("")
    _add_case_content(document, text)
    document.add_paragraph("")

def _group_cases_by_category(
    tcms,
    cases: list[dict[str, Any]],
) -> dict[str, list[dict[str, Any]]]:
    """
    Grupuje test case po kategorii.
    Jeśli case ma tylko ID kategorii, pobiera nazwy kategorii z Kiwi.
    """
    grouped: dict[str, list[dict[str, Any]]] = {}

    categories = tcms.exec.Category.filter({})
    category_map = {
        category["id"]: category["name"]
        for category in categories
    }

    for case in cases:
        category_id = case.get("category")
        category_name = category_map.get(category_id, "Bez kategorii")

        grouped.setdefault(str(category_name), []).append(case)

    return grouped

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_table_of_contents(document):
    """Dodaje spis treści do dokumentu Word."""
    paragraph = document.add_paragraph()

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Spis treści (zaktualizuj w Wordzie)"

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

def export_plan_to_docx(
    tcms,
    plan_name: str,
    plan_id: int,
    cases: list[dict[str, Any]],
) -> Path:
    document = Document()
    _set_doc_style(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"Test Plan: {plan_name}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(f"ID planu: {plan_id}")

    summary = document.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary.add_run(f"Liczba test case: {len(cases)}")

    document.add_paragraph("")

    toc_heading = document.add_paragraph(style="Heading 1")
    toc_heading.add_run("Spis treści")
    _add_table_of_contents(document)

    document.add_page_break()

    grouped_cases = _group_cases_by_category(tcms, cases)

    for category_name in sorted(grouped_cases.keys()):
        category_cases = grouped_cases[category_name]
        category_count = len(category_cases)

        category_heading = document.add_paragraph(style="Heading 1")
        category_heading.add_run(
            f"Kategoria: {category_name} ({category_count} Test Cases)"
        )

        for case in sorted(category_cases, key=lambda x: x.get("id", 0)):
            _add_case_section(document, case)

    output_dir = _ensure_output_dir()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", plan_name).strip("_")
    output_path = output_dir / f"{safe_name}_(ID_{plan_id})_{timestamp}.docx"

    document.save(output_path)
    return output_path

def export_product_to_docx(
    tcms,
    product_name: str,
    product_id: int,
    cases: list[dict[str, Any]],
    category_name: str | None = None,
) -> Path:
    document = Document()
    _set_doc_style(document)

    title_text = f"Test Cases: Projekt {product_name}"
    if category_name:
        title_text += f" / Kategoria {category_name}"

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_text = f"ID projektu: {product_id}"
    if category_name:
        subtitle_text += f" | Kategoria: {category_name}"
    subtitle.add_run(subtitle_text)

    summary = document.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary.add_run(f"Liczba test case: {len(cases)}")

    document.add_paragraph("")

    toc_heading = document.add_paragraph()
    toc_run = toc_heading.add_run("Spis treści")
    toc_run.bold = True
    toc_run.font.size = Pt(14)

    _add_table_of_contents(document)

    document.add_page_break()

    grouped_cases = _group_cases_by_category(tcms, cases)

    for current_category_name in sorted(grouped_cases.keys()):
        category_cases = grouped_cases[current_category_name]
        category_count = len(category_cases)

        category_heading = document.add_paragraph(style="Heading 1")
        category_heading.add_run(
            f"Kategoria: {current_category_name} ({category_count} Test Cases)"
        )

        for case in sorted(category_cases, key=lambda x: x.get("id", 0)):
            _add_case_section(document, case)

    output_dir = _ensure_output_dir()

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    safe_product_name = re.sub(r"[^a-zA-Z0-9]+", "_", product_name).strip("_")

    if category_name:
        safe_category_name = re.sub(r"[^a-zA-Z0-9]+", "_", category_name).strip("_")
        file_name = f"Test Cases - Projekt {safe_product_name} - Kategoria {safe_category_name}_{timestamp}.docx"
    else:
        file_name = f"Test Cases - Projekt {safe_product_name}_{timestamp}.docx"

    output_path = output_dir / file_name
    document.save(output_path)

    return output_path

def export_run_to_docx(
    tcms,
    run_name: str,
    run_id: int,
    executions: list[dict[str, Any]],
) -> Path:
    document = Document()
    _set_doc_style(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"Test Run: {run_name}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(f"ID Test Run: {run_id}")

    summary = document.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary.add_run(f"Liczba wykonań: {len(executions)}")

    document.add_paragraph("")

    toc_heading = document.add_paragraph()
    toc_run = toc_heading.add_run("Spis treści")
    toc_run.bold = True
    toc_run.font.size = Pt(14)

    _add_table_of_contents(document)
    document.add_page_break()

    statuses = tcms.exec.TestExecutionStatus.filter({})
    status_map = {
        status["id"]: status["name"]
        for status in statuses
    }

    for execution in sorted(executions, key=lambda x: x.get("id", 0)):
        _add_execution_section(document, tcms, execution, status_map)

    output_dir = _ensure_output_dir()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", run_name).strip("_")
    output_path = output_dir / f"Test_Run_{safe_name}_ID-{run_id}_{timestamp}.docx"

    document.save(output_path)
    return output_path


def _get_value_name(value) -> str:
    if isinstance(value, dict):
        return str(value.get("name") or value.get("value") or value.get("id") or "")
    return str(value or "")


def _add_execution_section(
    document: Document,
    tcms,
    execution: dict[str, Any],
    status_map: dict[int, str],
) -> None:
    case_id = execution.get("case")
    execution_id = execution.get("id")

    case = None

    # Pobieramy pełny Test Case.
    # W Twojej instancji TestCase.filter({"id": ...}) zwraca pełniejsze dane niż TestCase.get().
    if case_id:
        try:
            found_cases = tcms.exec.TestCase.filter({"id": case_id})
            if found_cases:
                case = found_cases[0]
        except Exception:
            case = None

    if case:
        case_summary = case.get("summary", "") or f"TC {case_id}"
        case_text = case.get("text", "")
    else:
        case_summary = execution.get("summary") or f"TC {case_id}"
        case_text = ""

    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(f"TC-{case_id}: {case_summary}")

    info_table = document.add_table(rows=7, cols=2)
    info_table.style = "Table Grid"
    info_table.autofit = False

    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(13.5)

    rows = info_table.rows

    rows[0].cells[0].text = "Execution ID"
    rows[0].cells[1].text = str(execution_id)

    rows[1].cells[0].text = "Test Case ID"
    rows[1].cells[1].text = str(case_id)

    rows[2].cells[0].text = "Nazwa testu"
    rows[2].cells[1].text = str(case_summary)

    status_id = execution.get("status")
    status_name = status_map.get(status_id, str(status_id))

    rows[3].cells[0].text = "Status"
    rows[3].cells[1].text = status_name

    rows[4].cells[0].text = "Tester"
    rows[4].cells[1].text = str(
        execution.get("tested_by__username")
        or execution.get("tested_by")
        or ""
    )

    rows[5].cells[0].text = "Data rozpoczęcia"
    rows[5].cells[1].text = str(execution.get("start_date") or "")

    rows[6].cells[0].text = "Data zakończenia"
    rows[6].cells[1].text = str(execution.get("stop_date") or "")

    actual_result = execution.get("actual_result") or execution.get("notes") or ""

    if actual_result:
        document.add_paragraph("Actual result", style="Heading 2")
        _add_case_content(document, actual_result)

    if case_text:
        document.add_paragraph("Treść test case", style="Heading 2")
        _add_case_content(document, case_text)

    document.add_paragraph("")

def _report_get_status_name(
        execution: dict[str, Any],
        status_map: dict[int, str],
) -> str:
    """
    Zwraca nazwę statusu wykonania testu.
    """
    status_id = execution.get("status")

    return str(
        execution.get("status__name")
        or execution.get("status_name")
        or status_map.get(status_id)
        or status_id
        or "BRAK STATUSU"
    )

def _report_get_tester_name(execution: dict[str, Any]) -> str:
    """
    Zwraca nazwę testera.
    """
    tester = (
            execution.get("tested_by__username")
            or execution.get("tested_by__email")
            or execution.get("tested_by")
            or ""
    )

    return _get_value_name(tester)

def _report_get_comment_text(comment: Any) -> str:
    """
    Pobiera treść komentarza niezależnie od formatu odpowiedzi Kiwi.
    """
    if comment is None:
        return ""

    if isinstance(comment, str):
        return clean_inline_formatting(comment)

    if isinstance(comment, dict):
        value = (
                comment.get("comment")
                or comment.get("text")
                or comment.get("body")
                or comment.get("content")
                or comment.get("description")
                or ""
        )

        return clean_inline_formatting(str(value))

    return clean_inline_formatting(str(comment))

def _report_get_comment_author(comment: Any) -> str:
    """
    Pobiera autora komentarza.
    """
    if not isinstance(comment, dict):
        return ""

    author = (
            comment.get("user__username")
            or comment.get("author__username")
            or comment.get("username")
            or comment.get("user")
            or comment.get("author")
            or comment.get("created_by")
            or ""
    )

    return _get_value_name(author)

def _report_get_comment_date(comment: Any) -> str:
    """
    Pobiera datę komentarza.
    """
    if not isinstance(comment, dict):
        return ""

    return str(
        comment.get("submit_date")
        or comment.get("created_at")
        or comment.get("created")
        or comment.get("date")
        or comment.get("timestamp")
        or ""
    )

def _report_add_comments(
        document: Document,
        comments: list[Any],
) -> None:
    """
    Dodaje do dokumentu komentarze wykonania testu.
    """
    document.add_paragraph(
        "Komentarze do wykonania",
        style="Heading 2",
    )

    if not comments:
        paragraph = document.add_paragraph("Brak komentarzy.")
        _set_paragraph_spacing(paragraph)
        return

    for index, comment in enumerate(comments, start=1):
        text = _report_get_comment_text(comment)
        author = _report_get_comment_author(comment)
        comment_date = _report_get_comment_date(comment)

        header_parts = []

        if author:
            header_parts.append(f"Autor: {author}")

        if comment_date:
            header_parts.append(f"Data: {comment_date}")

        header = " | ".join(header_parts)

        comment_table = document.add_table(rows=2, cols=2)
        comment_table.style = "Table Grid"
        comment_table.autofit = False

        for row in comment_table.rows:
            row.cells[0].width = Cm(3.5)
            row.cells[1].width = Cm(14)

        comment_table.rows[0].cells[0].text = "Komentarz"
        comment_table.rows[0].cells[1].text = str(index)

        comment_table.rows[1].cells[0].text = "Autor / data"
        comment_table.rows[1].cells[1].text = header or "Brak danych"

        if text:
            paragraph = document.add_paragraph()
            paragraph.add_run(text)
            _set_paragraph_spacing(
                paragraph,
                before=2,
                after=4,
            )
        else:
            paragraph = document.add_paragraph(
                "[Komentarz bez treści]"
            )
            _set_paragraph_spacing(
                paragraph,
                before=2,
                after=4,
            )

def _report_build_status_summary(
        executions: list[dict[str, Any]],
        status_map: dict[int, str],
) -> dict[str, int]:
    """
    Zlicza wykonania według statusu.
    """
    result: dict[str, int] = {}

    for execution in executions:
        status_name = _report_get_status_name(
            execution,
            status_map,
        )

        result[status_name] = result.get(status_name, 0) + 1

    return result

def _report_add_status_summary(
        document: Document,
        executions: list[dict[str, Any]],
        status_map: dict[int, str],
) -> None:
    """
    Dodaje tabelę podsumowania statusów.
    """
    document.add_paragraph(
        "Podsumowanie realizacji Test Run",
        style="Heading 1",
    )

    status_summary = _report_build_status_summary(
        executions,
        status_map,
    )

    total = len(executions)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False

    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Liczba"
    table.rows[0].cells[2].text = "Udział"

    for row in table.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(3)
        row.cells[2].width = Cm(4)

    sorted_statuses = sorted(
        status_summary.items(),
        key=lambda item: (-item[1], item[0]),
    )

    for status_name, count in sorted_statuses:
        row = table.add_row().cells

        percentage = (
            count / total * 100
            if total
            else 0
        )

        row[0].text = status_name
        row[1].text = str(count)
        row[2].text = f"{percentage:.1f}%".replace(".", ",")

    total_row = table.add_row().cells
    total_row[0].text = "RAZEM"
    total_row[1].text = str(total)
    total_row[2].text = "100,0%" if total else "0,0%"

    document.add_paragraph("")

def _report_add_execution_list(
        document: Document,
        executions: list[dict[str, Any]],
        status_map: dict[int, str],
) -> None:
    """
    Dodaje zbiorczą tabelę wszystkich wykonań.
    """
    document.add_paragraph(
        "Zestawienie wyników testów",
        style="Heading 1",
    )

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.autofit = False

    headers = [
        "Lp.",
        "Test Case",
        "Nazwa testu",
        "Status",
        "Tester",
        "Data zakończenia",
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    widths = [1, 2.2, 6.8, 2.8, 2.5, 2.8]

    for column_index, width in enumerate(widths):
        _set_column_width(
            table,
            column_index,
            width,
        )

    sorted_executions = sorted(
        executions,
        key=lambda item: item.get("id", 0),
    )

    for number, execution in enumerate(
            sorted_executions,
            start=1,
    ):
        row = table.add_row().cells

        case_id = execution.get("case", "")
        summary = (
                execution.get("case__summary")
                or execution.get("summary")
                or ""
        )

        row[0].text = str(number)
        row[1].text = f"TC-{case_id}"
        row[2].text = str(summary)
        row[3].text = _report_get_status_name(
            execution,
            status_map,
        )
        row[4].text = _report_get_tester_name(execution)
        row[5].text = str(
            execution.get("stop_date")
            or ""
        )

    document.add_paragraph("")

def _report_get_case(
        tcms,
        case_id: int | None,
) -> dict[str, Any] | None:
    """
    Pobiera pełny Test Case.
    """
    if not case_id:
        return None

    try:
        found_cases = tcms.exec.TestCase.filter({
            "id": case_id
        })

        if found_cases:
            return found_cases[0]

    except Exception as error:
        print(
            f"Nie udało się pobrać Test Case ID "
            f"{case_id}: {error}"
        )

    return None

def _report_add_execution_section(
        document: Document,
        tcms,
        execution: dict[str, Any],
        status_map: dict[int, str],
) -> None:
    """
    Dodaje rozszerzone informacje o wykonaniu testu.
    """
    case_id = execution.get("case")
    execution_id = execution.get("id")

    case = _report_get_case(
        tcms,
        case_id,
    )

    if case:
        case_summary = (
                case.get("summary")
                or f"TC {case_id}"
        )
        case_text = case.get("text") or ""
    else:
        case_summary = (
                execution.get("case__summary")
                or execution.get("summary")
                or f"TC {case_id}"
        )
        case_text = ""

    heading = document.add_paragraph(
        style="Heading 1"
    )
    heading.add_run(
        f"TC-{case_id}: {case_summary}"
    )

    info_table = document.add_table(
        rows=8,
        cols=2,
    )
    info_table.style = "Table Grid"
    info_table.autofit = False

    for row in info_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(13.5)

    values = [
        ("Execution ID", execution_id),
        ("Test Case ID", case_id),
        ("Nazwa testu", case_summary),
        (
            "Status",
            _report_get_status_name(
                execution,
                status_map,
            ),
        ),
        (
            "Tester",
            _report_get_tester_name(execution),
        ),
        (
            "Data rozpoczęcia",
            execution.get("start_date") or "",
        ),
        (
            "Data zakończenia",
            execution.get("stop_date") or "",
        ),
        (
            "Liczba komentarzy",
            len(execution.get("comments") or []),
        ),
    ]

    for row_index, (label, value) in enumerate(values):
        info_table.rows[row_index].cells[0].text = str(label)
        info_table.rows[row_index].cells[1].text = str(
            value or ""
        )

    actual_result = (
            execution.get("actual_result")
            or execution.get("notes")
            or ""
    )

    document.add_paragraph(
        "Wynik wykonania",
        style="Heading 2",
    )

    if actual_result:
        _add_case_content(
            document,
            str(actual_result),
        )
    else:
        paragraph = document.add_paragraph(
            "Brak opisu wyniku wykonania."
        )
        _set_paragraph_spacing(paragraph)

    comments = execution.get("comments") or []

    _report_add_comments(
        document,
        comments,
    )

    if case_text:
        document.add_paragraph(
            "Treść Test Case",
            style="Heading 2",
        )

        _add_case_content(
            document,
            case_text,
        )

    document.add_paragraph("")

def export_run_report_to_docx(
    tcms,
    run_name: str,
    run_id: int,
    executions: list[dict[str, Any]],
    report_title: str | None = None,
    file_prefix: str = "Raport_Test_Run",
) -> Path:
    """
    Generuje rozszerzony raport wykonania Test Run.

    Funkcja jest niezależna od export_run_to_docx
    i nie zmienia dotychczasowego formatu eksportu.
    """
    document = Document()
    _set_doc_style(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if report_title is None:
        report_title = f"Raport z realizacji Test Run: {run_name}"

    title_run = title.add_run(report_title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(
        f"ID Test Run: {run_id}"
    )

    execution_count = document.add_paragraph()
    execution_count.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    execution_count.add_run(
        f"Liczba wykonań: {len(executions)}"
    )

    generation_date = document.add_paragraph()
    generation_date.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    generation_date.add_run(
        "Data wygenerowania raportu: "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    document.add_paragraph("")

    statuses = tcms.exec.TestExecutionStatus.filter({})

    status_map = {
        status["id"]: status["name"]
        for status in statuses
    }

    _report_add_status_summary(
        document,
        executions,
        status_map,
    )

    _report_add_execution_list(
        document,
        executions,
        status_map,
    )

    toc_heading = document.add_paragraph()
    toc_run = toc_heading.add_run(
        "Spis treści"
    )
    toc_run.bold = True
    toc_run.font.size = Pt(14)

    _add_table_of_contents(document)

    document.add_page_break()

    for execution in sorted(
            executions,
            key=lambda item: item.get("id", 0),
    ):
        _report_add_execution_section(
            document=document,
            tcms=tcms,
            execution=execution,
            status_map=status_map,
        )

    output_dir = _ensure_output_dir()
    timestamp = datetime.now().strftime(
        "%Y-%m-%d_%H-%M-%S"
    )

    safe_name = re.sub(
        r"[^a-zA-Z0-9]+",
        "_",
        run_name,
    ).strip("_")

    output_path = output_dir / (
        f"{file_prefix}_{safe_name}_"
        f"ID-{run_id}_{timestamp}.docx"
    )

    document.save(output_path)

    return output_path